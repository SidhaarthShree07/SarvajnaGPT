from __future__ import annotations

import os
from typing import Optional

from fastapi import APIRouter, HTTPException
from pydantic import BaseModel, Field

try:
    from docx import Document  # python-docx
except Exception as _e:
    Document = None  # type: ignore


router = APIRouter(prefix="/api/word", tags=["word"])

REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), os.pardir))
AGENT_BASE_DIR = os.path.join(REPO_ROOT, "agent_output")
os.makedirs(AGENT_BASE_DIR, exist_ok=True)


def _resolve_under_base(path_like: str) -> str:
    p = path_like
    if not os.path.isabs(p):
        p = os.path.join(AGENT_BASE_DIR, p)
    p = os.path.abspath(p)
    base = os.path.abspath(AGENT_BASE_DIR)
    if os.path.commonpath([p, base]) != base:
        raise ValueError("Target path outside agent base directory")
    return p


class WordPreviewRequest(BaseModel):
    target_rel: str = Field(..., description="Relative path to .docx under agent_output, e.g., docs/demo.docx")
    paragraph: str = Field("", description="Paragraph text to insert")


@router.post("/preview")
def preview(req: WordPreviewRequest) -> dict:
    target_abs = _resolve_under_base(req.target_rel)
    if not req.target_rel.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="target_rel must end with .docx")
    return {
        "summary": f"Create Word document and insert 1 paragraph ({len(req.paragraph)} chars)",
        "target_rel": req.target_rel,
        "target_abs": target_abs,
    }


@router.post("/execute")
def execute(req: WordPreviewRequest) -> dict:
    if Document is None:
        raise HTTPException(status_code=500, detail="python-docx is not available; install it from backend/requirements.txt")
    prev = preview(req)
    target_abs = prev["target_abs"]
    os.makedirs(os.path.dirname(target_abs), exist_ok=True)
    try:
        doc = Document()
        if req.paragraph:
            doc.add_paragraph(req.paragraph)
        else:
            doc.add_paragraph("")
        doc.save(target_abs)
        return {"created": True, "path": target_abs}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
