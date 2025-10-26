"""CUA-only Power Router (pure)

This router provides a minimal shape-compatible subset of the existing power router
implemented only with the CUA adapter. It aims to:
- Provide a direct "open document" endpoint compatible with the caller but without
  any Win32/COM/UIAutomation logic.
- Optionally trigger CUA Snap Assist tile selection for Word or general windows.

Important limitations:
- No VS Code or filesystem orchestration here; this is purely a sketch of CUA-only
  behavior for opening documents and hinting a snap tile selection.
- It relies entirely on `os.startfile` and the CUA adapter’s token selection.
"""
from __future__ import annotations

from typing import Any, Dict, Optional
from pathlib import Path
from fastapi import APIRouter, HTTPException
from pydantic import BaseModel, Field
import os
import hashlib
import datetime as _dt
import re as _re
import sqlite3 as _sqlite3
import html as _html
router = APIRouter(prefix="/api/power", tags=["power-cua-only"])

# Soft import CUA adapter
try:
    from .cua_adapter import (
        cua_runtime_status,
        select_snap_assist_tile,
        snap_current_and_select,
        trigger_snap,
        capture_inline_selection,
        capture_full_document_text_and_restore_selection,
        open_path as cua_open_path,
        open_path_background as cua_open_path_background,
        open_vscode as cua_open_vscode,
        open_browser_to_path as cua_open_browser_to_path,
        wait_for_focus,
        focus_previous_window,
        focus_window_by_tokens,
        wait_for_window_appearance,
        ensure_focus,
        ensure_focus_top,
        get_focused_window_name,
        paste_text_to_foreground_app,
        paste_rich_text_to_foreground_app,
        arrange_three_columns,
        arrange_right_stack,
        layout_left_right_stack,
    )  # type: ignore
except Exception:  # pragma: no cover
    try:
        from cua_adapter import (
            cua_runtime_status,
            select_snap_assist_tile,
            snap_current_and_select,
            trigger_snap,
            capture_inline_selection,
            capture_full_document_text_and_restore_selection,
            open_path as cua_open_path,
            open_path_background as cua_open_path_background,
            open_vscode as cua_open_vscode,
            open_browser_to_path as cua_open_browser_to_path,
            wait_for_focus,
            focus_previous_window,
            focus_window_by_tokens,
            wait_for_window_appearance,
            ensure_focus,
            ensure_focus_top,
            get_focused_window_name,
            paste_text_to_foreground_app,
            paste_rich_text_to_foreground_app,
            arrange_three_columns,
            layout_left_right_stack,
        )  # type: ignore
    except Exception:
        cua_runtime_status = None  # type: ignore
        select_snap_assist_tile = None  # type: ignore


class OpenDocCUARequest(BaseModel):
    path: str = Field(..., description="Absolute or user-expanded path to a document (e.g., .docx)")
    snap: bool | None = Field(True, description="If True, attempt a CUA Snap Assist selection for Word")


# ----------------------------- Text cleaning utils -----------------------------
def _strip_think_blocks(s: str) -> str:
    """Remove <think>...</think> blocks (case-insensitive, multiline)."""
    try:
        return _re.sub(r"(?is)<think[^>]*>.*?</think>", "", s or "").strip()
    except Exception:
        return s or ""


def _plain_from_markdown(md: str) -> str:
    """Best-effort conversion from Markdown to plain text suitable for pasting into Word.
    - Removes headings markers, code fences, backticks, emphasis markers, and <think> blocks.
    - Keeps bullet leading markers ('- ' and '* ') and numbering (e.g., '1. ').
    """
    s = _strip_think_blocks(md or "")
    lines = (s or "").splitlines()
    out_lines: list[str] = []
    in_code = False
    for line in lines:
        l = line.rstrip("\n\r")
        if l.strip().startswith("```"):
            in_code = not in_code
            continue  # drop fence lines
        if not in_code:
            # Strip heading markers at start
            l = _re.sub(r"^\s*#{1,6}\s*", "", l)
            # Remove bold/italic markers
            l = l.replace("**", "").replace("__", "")
            l = l.replace("`", "")
            stripped = l.lstrip()
            if stripped.startswith("* "):
                # keep leading '* ' as bullet
                # remove other stray '*' in the line
                prefix_len = len(l) - len(stripped)
                l = l[:prefix_len] + "* " + stripped[2:].replace("*", "")
            else:
                # remove stray '*' used for italics
                l = l.replace("*", "")
            # remove stray '_' used for italics
            l = l.replace("_", "")
        out_lines.append(l)
    return ("\n".join(out_lines)).strip()


def _docx_from_markdown(md: str, out_path: str) -> bool:
    """Create a .docx file from Markdown with minimal formatting (headings, bullets, bold/italic).
    Returns True on success. Falls back to plain text if python-docx isn't available.
    """
    try:
        import docx  # type: ignore
    except Exception:
        # Fallback: write plain text
        try:
            Path(out_path).write_text(_plain_from_markdown(md), encoding='utf-8')
            return True
        except Exception:
            return False
    try:
        doc = docx.Document()

        def add_formatted_paragraph(text: str, style: str | None = None):
            p = doc.add_paragraph()
            if style:
                try:
                    p.style = style
                except Exception:
                    pass
            # Inline bold/italic parsing
            # Matches: **bold**, __bold__, *italic*, _italic_
            pattern = _re.compile(r"(\*\*[^*]+\*\*|__[^_]+__|\*[^*]+\*|_[^_]+_)")
            idx = 0
            for m in pattern.finditer(text):
                if m.start() > idx:
                    p.add_run(text[idx:m.start()])
                chunk = m.group(0)
                run_text = chunk
                run_bold = False
                run_italic = False
                if chunk.startswith("**") and chunk.endswith("**"):
                    run_text = chunk[2:-2]
                    run_bold = True
                elif chunk.startswith("__") and chunk.endswith("__"):
                    run_text = chunk[2:-2]
                    run_bold = True
                elif chunk.startswith("*") and chunk.endswith("*"):
                    run_text = chunk[1:-1]
                    run_italic = True
                elif chunk.startswith("_") and chunk.endswith("_"):
                    run_text = chunk[1:-1]
                    run_italic = True
                r = p.add_run(run_text)
                if run_bold:
                    r.bold = True
                if run_italic:
                    r.italic = True
                idx = m.end()
            if idx < len(text):
                p.add_run(text[idx:])

        s = _strip_think_blocks(md or "")
        lines = s.splitlines()
        in_code = False
        for raw in lines:
            line = raw.rstrip("\n\r")
            if line.strip().startswith("```"):
                in_code = not in_code
                continue
            if in_code:
                # Code block: add as plain paragraph
                doc.add_paragraph(line)
                continue
            # Headings
            if line.lstrip().startswith("### "):
                add_formatted_paragraph(line.lstrip()[4:], style='Heading 3')
                continue
            if line.lstrip().startswith("## "):
                add_formatted_paragraph(line.lstrip()[3:], style='Heading 2')
                continue
            if line.lstrip().startswith("# "):
                add_formatted_paragraph(line.lstrip()[2:], style='Heading 1')
                continue
            # Lists
            lstr = line.lstrip()
            if lstr.startswith("- ") or lstr.startswith("* "):
                try:
                    add_formatted_paragraph(lstr[2:], style='List Bullet')
                except Exception:
                    doc.add_paragraph(lstr[2:])
                continue
            if _re.match(r"^\s*\d+[\.)]\s+", line):
                item = _re.sub(r"^\s*\d+[\.)]\s+", "", line)
                try:
                    add_formatted_paragraph(item, style='List Number')
                except Exception:
                    doc.add_paragraph(item)
                continue
            # Paragraph
            add_formatted_paragraph(line)

        doc.save(out_path)
        return True
    except Exception:
        try:
            Path(out_path).write_text(_plain_from_markdown(md), encoding='utf-8')
            return True
        except Exception:
            return False


def _rtf_from_markdown(md: str) -> str:
    """Convert minimal Markdown to simple RTF supporting bold/italic and basic paragraphs/lists.
    This is not exhaustive but handles common inline formatting like **bold** and *italic*.
    """
    def esc(s: str) -> str:
        s = s.replace('\\', r'\\').replace('{', r'\{').replace('}', r'\}')
        # Replace newlines with RTF paragraph breaks
        return s
    s = _strip_think_blocks(md or '')
    lines = s.splitlines()
    parts = [r'{\rtf1\ansi\deff0']
    in_code = False
    for raw in lines:
        line = raw.rstrip('\r\n')
        if line.strip().startswith('```'):
            in_code = not in_code
            # Emit a paragraph break on fence toggle
            parts.append('\\par ')
            continue
        txt = esc(line)
        # Inline bold/italic
        # Process bold first, then italic, using a scan loop to avoid overlap issues
        def fmt_inline(s0: str) -> str:
            out = ''
            i = 0
            while i < len(s0):
                if s0.startswith('**', i):
                    j = s0.find('**', i + 2)
                    if j != -1:
                        content = s0[i+2:j]
                        out += r'\b ' + esc(content) + r'\b0'
                        i = j + 2
                        continue
                if s0.startswith('__', i):
                    j = s0.find('__', i + 2)
                    if j != -1:
                        content = s0[i+2:j]
                        out += r'\b ' + esc(content) + r'\b0'
                        i = j + 2
                        continue
                if s0.startswith('*', i):
                    j = s0.find('*', i + 1)
                    if j != -1:
                        content = s0[i+1:j]
                        out += r'\i ' + esc(content) + r'\i0'
                        i = j + 1
                        continue
                if s0.startswith('_', i):
                    j = s0.find('_', i + 1)
                    if j != -1:
                        content = s0[i+1:j]
                        out += r'\i ' + esc(content) + r'\i0'
                        i = j + 1
                        continue
                out += s0[i]
                i += 1
            return out
        txt = fmt_inline(txt)
        if _re.match(r'^\s*[-*]\s+', line):
            content = _re.sub(r'^\s*[-*]\s+', '', line)
            parts.append(r'\bullet\tab ' + fmt_inline(esc(content)) + r'\par ')
        elif _re.match(r'^\s*\d+[\.)]\s+', line):
            content = _re.sub(r'^\s*\d+[\.)]\s+', '', line)
            parts.append(fmt_inline(esc(content)) + r'\par ')
        else:
            parts.append(txt + r'\par ')
    parts.append('}')
    return ''.join(parts)


def _guess_code_language(text: str) -> str:
    lt = (text or '').lower()
    if '```html' in lt or '<html' in lt or 'doctype html' in lt or 'html>' in lt:
        return 'html'
    if '```python' in lt or 'def ' in lt and 'import ' in lt:
        return 'python'
    if '```typescript' in lt or 'typescript' in lt:
        return 'typescript'
    if '```javascript' in lt or '```js' in lt or '<script' in lt:
        return 'javascript'
    return 'plain'


def _extract_code_from_llm(text: str, lang_hint: Optional[str] = None) -> str:
    """Extract code-only content from LLM output.
    - Prefer fenced code blocks for the detected or hinted language.
    - For HTML: if <html>...</html> is present, extract that region.
    - Strip <think> blocks and markdown prose.
    """
    s = _strip_think_blocks(text or '')
    lhint = (lang_hint or '').lower() if lang_hint else None
    # 1) Try fenced code blocks
    try:
        fence_re = _re.compile(r"```([a-zA-Z0-9]+)?\n([\s\S]*?)```", _re.MULTILINE)
        matches = list(fence_re.finditer(s))
        if matches:
            # If language hinted, pick that; else pick the longest block
            best = None
            best_len = -1
            for m in matches:
                lang = (m.group(1) or '').lower()
                body = m.group(2) or ''
                if lhint and lang and lhint in lang:
                    return body.strip()
                if len(body) > best_len:
                    best = body
                    best_len = len(body)
            if best is not None:
                return best.strip()
    except Exception:
        pass
    # 2) HTML region extraction
    if (lhint == 'html') or ('<html' in s.lower() and '</html>' in s.lower()):
        try:
            start = s.lower().find('<html')
            end = s.lower().rfind('</html>')
            if start != -1 and end != -1:
                return s[start:end+7].strip()
        except Exception:
            pass
    # 3) Fallback: remove markdown headings / notes, keep code-looking lines
    lines = s.splitlines()
    out = []
    for line in lines:
        if line.strip().startswith('```'):
            continue
        if line.strip().startswith(('#', '###', '####', 'Notes:', 'Note:', '>')):
            continue
        # Skip list bullets unless it's HTML tags
        if line.strip().startswith(('- ', '* ')) and not line.strip().startswith(('<', '</')):
            continue
        out.append(line)
    return '\n'.join(out).strip()


def _html_title_from_file(path: Path) -> Optional[str]:
    """Best-effort extraction of the <title>...</title> from an HTML file.
    Returns a trimmed, HTML-unescaped title string or None.
    """
    try:
        # Read the first 128KB which should include <head>
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            chunk = f.read(131072)
        m = _re.search(r"(?is)<title[^>]*>(.*?)</title>", chunk)
        if not m:
            return None
        t = m.group(1) or ''
        t = _html.unescape(t.strip())
        # Normalize internal whitespace
        t = _re.sub(r"\s+", " ", t)
        return t[:256] if t else None
    except Exception:
        return None


# ---------------------------------------------------------------------------
# Compatibility shim for main.py direct endpoint
# ---------------------------------------------------------------------------
class OpenDocIntelligentlyRequest(BaseModel):
    abs_path: str = Field(..., description="Absolute path to a document to open")
    request_id: Optional[str] = Field(None, description="Optional idempotency key")
    split_screen: bool | None = Field(True, description="Attempt CUA snap assist selection")
    side: Optional[str] = Field('right', description="Desired side for current window before selection (hint only)")
    snap: bool | None = Field(True, description="Alias for split_screen; when True, attempt selection")


def compute_open_doc_signature(req: 'OpenDocIntelligentlyRequest') -> str:
    base = f"open_doc:{os.path.abspath(req.abs_path)}"
    return hashlib.md5(base.encode('utf-8')).hexdigest()


def open_doc_intelligently(req: 'OpenDocIntelligentlyRequest') -> Dict[str, Any]:
    """CUA-only implementation used by main.py direct endpoint.

    - Launches the document via cua_open_path
    - If split/snap requested, triggers snap+tile selection using tokens derived from the file name
    Returns a result compatible with prior callers: includes 'opened', 'path', 'split_screen', and CUA flags.
    """
    p = Path(req.abs_path).expanduser()
    if not p.is_file():
        raise HTTPException(status_code=400, detail="abs_path not found")
    # Determine file type
    suffix = p.suffix.lower()
    is_code_like = suffix in {".py", ".js", ".ts", ".tsx", ".jsx", ".html", ".css"}
    is_word_like = suffix in {".doc", ".docx", ".rtf"}
    # Launch: Word in background; Code via VS Code in a NEW window; others normal
    if is_word_like and callable(cua_open_path_background):
        r_launch = cua_open_path_background(str(p))  # type: ignore[misc]
        if not r_launch.get("ok") and callable(cua_open_path):
            r_launch = cua_open_path(str(p))  # fallback
    elif is_code_like and callable(cua_open_vscode):
        try:
            r_launch = cua_open_vscode(str(p), True)  # type: ignore[misc]
        except Exception:
            r_launch = {"ok": False, "error": "vscode_launch_failed"}
        if not r_launch.get("ok") and callable(cua_open_path):
            r_launch = cua_open_path(str(p))
    else:
        r_launch = cua_open_path(str(p)) if callable(cua_open_path) else {"ok": False, "error": "helper_missing"}
    launched = bool(r_launch.get("ok"))

    # If HTML: open a browser preview in a separate window to enable reliable selection
    browser_info = None
    tri_snap = None
    if launched and suffix == '.html' and callable(cua_open_browser_to_path):
        try:
            browser_info = cua_open_browser_to_path(str(p), new_window=True)
        except Exception:
            browser_info = {"ok": False, "error": "browser_launch_failed"}
        # Wait briefly for the preview window to appear so selection can target it by file stem
        try:
            stem = p.stem
            html_title = _html_title_from_file(p)
            stem_space = stem.replace("-", " ")
            if callable(wait_for_window_appearance):
                wait_tokens = [p.name, stem, stem_space, p.stem, "microsoft edge", "edge", "google chrome", "chrome", "mozilla firefox", "firefox", "brave"]
                if html_title:
                    wait_tokens = [html_title] + wait_tokens
                _ = wait_for_window_appearance(wait_tokens, timeout_ms=7000)
        except Exception:
            pass
        # Pre-warm: ensure the preview window is in the recent MRU by briefly focusing it, then return to our app
        try:
            if callable(focus_window_by_tokens):
                focus_tokens = [stem, stem_space, "microsoft edge", "edge", "google chrome", "chrome", "mozilla firefox", "firefox", "brave"]
                if html_title:
                    focus_tokens = [html_title] + focus_tokens
                focus_window_by_tokens(focus_tokens)  # bring preview forward once
                # return focus to our app window before snapping
                focus_window_by_tokens(["sarvajña", "sarvajna"])  # best-effort
        except Exception:
            pass
    # Attempt selection based on request
    want_snap = bool(req.split_screen if req.split_screen is not None else req.snap)
    snap = None
    if want_snap:
        # Prefer exact document stem for Word selection; keep search minimal.
        stem = p.stem
        tokens = [stem]
        # For Word files, only add a single Word hint. Avoid broad, noisy tokens.
        if is_word_like and not is_code_like:
            tokens.append("microsoft word")
        # For Code: wait for VS Code window to appear and focus it (so we can snap Code LEFT)
        if is_code_like and suffix != '.html':
            try:
                if callable(wait_for_window_appearance):
                    _ = wait_for_window_appearance([p.name, p.stem, "visual studio code", "code"], timeout_ms=2000)
                else:
                    _ = None
            except Exception:
                _ = None
            try:
                if callable(focus_window_by_tokens):
                    focus_window_by_tokens([p.name, p.stem, "visual studio code", "code"])  # best-effort focus on Code
                elif callable(wait_for_focus):
                    wait_for_focus([p.name, p.stem, "visual studio code", "code"], timeout_ms=800)
            except Exception:
                pass
            try:
                import time as _t
                _t.sleep(0.08)
            except Exception:
                pass
        # For Word: wait for Word to appear, then focus Word so we can snap it LEFT
        if is_word_like and not is_code_like:
            try:
                if callable(wait_for_window_appearance):
                    _ = wait_for_window_appearance([stem, "microsoft word", "word"], timeout_ms=2000)
                else:
                    _ = None
            except Exception:
                _ = None
            # Focus the Word window explicitly
            try:
                if callable(focus_window_by_tokens):
                    focus_window_by_tokens([stem, "microsoft word", "word"])  # best-effort
                elif callable(wait_for_focus):
                    wait_for_focus([stem, "microsoft word", "word"], timeout_ms=800)
            except Exception:
                pass
            try:
                import time as _t
                _t.sleep(0.08)
            except Exception:
                pass
        # Ensure Word remains focused for the snap (skip re-focusing our app)
        if is_word_like and not is_code_like:
            try:
                if callable(focus_window_by_tokens):
                    focus_window_by_tokens([stem, "microsoft word", "word"])  # re-affirm focus on Word
            except Exception:
                pass
        # Try snap+select; if grid not ready, retry once after a brief pause
        if callable(snap_current_and_select):  # type: ignore[truthy-bool]
            # For Word: snap Word LEFT and select SarvajñaGPT browser on RIGHT; For Code: snap Code LEFT and select SarvajñaGPT browser on RIGHT; else, default
            is_word_like = is_word_like
            if (not is_code_like) and is_word_like:
                side = 'left'
                # Prefer selecting the SarvajñaGPT tab specifically on the RIGHT. Avoid generic browser-only fallbacks.
                # Include common title variations and normalized forms to improve matching reliability.
                attempt_sets: list[list[str]] = [
                    ["sarvajnagpt", "microsoft edge"], ["sarvajna gpt", "microsoft edge"],
                    ["sarvajña", "microsoft edge"], ["sarvajna", "microsoft edge"],
                    ["sarvajnagpt", "google chrome"], ["sarvajna gpt", "google chrome"],
                    ["sarvajña", "google chrome"], ["sarvajna", "google chrome"],
                    ["sarvajnagpt", "mozilla firefox"], ["sarvajna gpt", "mozilla firefox"],
                    ["sarvajña", "mozilla firefox"], ["sarvajna", "mozilla firefox"],
                    ["sarvajnagpt", "brave"], ["sarvajna gpt", "brave"],
                    ["sarvajña", "brave"], ["sarvajna", "brave"],
                ]
                # Fallback: generic browser-only tokens to at least complete the split
                attempt_sets_generic: list[list[str]] = [
                    ["microsoft edge"], ["google chrome"], ["mozilla firefox"], ["brave"],
                ]
                import time as _tlim
                snap = {"attempted": True, "selected": False}
                start = _tlim.time()
                # Trigger snap once using the first tokens
                snap = snap_current_and_select(attempt_sets[0], snap_side=side)  # type: ignore[misc]
                if not snap.get("selected"):
                    # Within the remaining 2s window, try alternative tokens WITHOUT re-triggering snap
                    for toks in attempt_sets[1:]:
                        if (_tlim.time() - start) > 3.0:
                            break
                        try:
                            ok = bool(select_snap_assist_tile(toks)) if callable(select_snap_assist_tile) else False
                            if ok:
                                snap["selected"] = True
                                snap["tokens"] = toks
                                break
                        except Exception:
                            break
                # Diagnostics-guided retry: if a tile containing SarvajñaGPT was seen, target it explicitly
                if not snap.get("selected") and isinstance(snap.get("diagnostics"), dict):
                    try:
                        diag = snap.get("diagnostics") or {}
                        names = [str(n) for n in (diag.get("unique_names") or [])]
                        sar_cands = [n for n in names if ("sarvajna" in n.lower()) or ("sarvajña" in n.lower()) or ("sarvajnagpt" in n.lower())]
                        if sar_cands and callable(select_snap_assist_tile):
                            for cand in sar_cands[:4]:  # try up to 4 distinct candidates
                                try:
                                    # Try exact-name only, then with browser brand hints
                                    if bool(select_snap_assist_tile([cand])):
                                        snap["selected"] = True; snap["tokens"] = [cand]
                                        break
                                    for brand in ("microsoft edge", "google chrome", "mozilla firefox", "brave"):
                                        if bool(select_snap_assist_tile([cand, brand])):
                                            snap["selected"] = True; snap["tokens"] = [cand, brand]
                                            break
                                    if snap.get("selected"):
                                        break
                                except Exception:
                                    continue
                    except Exception:
                        pass
                # If still not selected, re-trigger once with generic tokens and allow brief follow-ups
                if not snap.get("selected"):
                    try:
                        snap = snap_current_and_select(attempt_sets_generic[0], snap_side=side)  # type: ignore[misc]
                    except Exception:
                        pass
                    if not snap.get("selected"):
                        g_start = _tlim.time()
                        for toks in attempt_sets_generic[1:]:
                            if (_tlim.time() - g_start) > 2.0:
                                break
                            try:
                                ok = bool(select_snap_assist_tile(toks)) if callable(select_snap_assist_tile) else False
                                if ok:
                                    snap["selected"] = True
                                    snap["tokens"] = toks
                                    break
                            except Exception:
                                break
            elif is_code_like and suffix != '.html':
                side = 'left'
                # Prefer SarvajñaGPT tab for right-side selection in double-split with VS Code
                attempt_sets: list[list[str]] = [
                    ["sarvajnagpt", "microsoft edge"], ["sarvajna gpt", "microsoft edge"],
                    ["sarvajña", "microsoft edge"], ["sarvajna", "microsoft edge"],
                    ["sarvajnagpt", "google chrome"], ["sarvajna gpt", "google chrome"],
                    ["sarvajña", "google chrome"], ["sarvajna", "google chrome"],
                    ["sarvajnagpt", "mozilla firefox"], ["sarvajna gpt", "mozilla firefox"],
                    ["sarvajña", "mozilla firefox"], ["sarvajna", "mozilla firefox"],
                    ["sarvajnagpt", "brave"], ["sarvajna gpt", "brave"],
                    ["sarvajña", "brave"], ["sarvajna", "brave"],
                ]
                # Fallback: generic browser tokens to ensure split completes if SarvajñaGPT not visible
                attempt_sets_generic: list[list[str]] = [["microsoft edge"], ["google chrome"], ["mozilla firefox"], ["brave"]]
                import time as _tlim
                snap = {"attempted": True, "selected": False}
                start = _tlim.time()
                # Trigger snap once using the first tokens
                snap = snap_current_and_select(attempt_sets[0], snap_side=side)  # type: ignore[misc]
                if not snap.get("selected"):
                    # Within the remaining window, try alternative tokens WITHOUT re-triggering snap
                    for toks in attempt_sets[1:]:
                        if (_tlim.time() - start) > 2.2:
                            break
                        try:
                            ok = bool(select_snap_assist_tile(toks)) if callable(select_snap_assist_tile) else False
                            if ok:
                                snap["selected"] = True
                                snap["tokens"] = toks
                                break
                        except Exception:
                            break
                # Diagnostics-guided retry: explicitly target any visible SarvajñaGPT tiles
                if not snap.get("selected") and isinstance(snap.get("diagnostics"), dict):
                    try:
                        diag = snap.get("diagnostics") or {}
                        names = [str(n) for n in (diag.get("unique_names") or [])]
                        sar_cands = [n for n in names if ("sarvajna" in n.lower()) or ("sarvajña" in n.lower()) or ("sarvajnagpt" in n.lower())]
                        if sar_cands and callable(select_snap_assist_tile):
                            for cand in sar_cands[:4]:
                                try:
                                    if bool(select_snap_assist_tile([cand])):
                                        snap["selected"] = True; snap["tokens"] = [cand]
                                        break
                                    for brand in ("microsoft edge", "google chrome", "mozilla firefox", "brave"):
                                        if bool(select_snap_assist_tile([cand, brand])):
                                            snap["selected"] = True; snap["tokens"] = [cand, brand]
                                            break
                                    if snap.get("selected"):
                                        break
                                except Exception:
                                    continue
                    except Exception:
                        pass
                # Final generic fallback: re-trigger and try generic brand tokens
                if not snap.get("selected"):
                    try:
                        snap = snap_current_and_select(attempt_sets_generic[0], snap_side=side)  # type: ignore[misc]
                    except Exception:
                        pass
                    if not snap.get("selected"):
                        g_start = _tlim.time()
                        for toks in attempt_sets_generic[1:]:
                            if (_tlim.time() - g_start) > 2.0:
                                break
                            try:
                                ok = bool(select_snap_assist_tile(toks)) if callable(select_snap_assist_tile) else False
                                if ok:
                                    snap["selected"] = True
                                    snap["tokens"] = toks
                                    break
                            except Exception:
                                break
            elif is_code_like and suffix == '.html':
                # New generic, order-agnostic flow per user request:
                # 1) Ensure browser+code windows exist, then perform right snap and select tiles without predetermined roles.
                try:
                    from .cua_adapter import layout_any_right_stack  # type: ignore
                except Exception:
                    try:
                        from cua_adapter import layout_any_right_stack  # type: ignore
                    except Exception:
                        layout_any_right_stack = None  # type: ignore
                try:
                    stem_space_arr = p.stem.replace("-", " ")
                    page_title = html_title if 'html_title' in locals() else _html_title_from_file(p)
                    code_title1 = f"{p.name} - Visual Studio Code"; code_title2 = f"{p.stem} - Visual Studio Code"
                    edge_title1 = f"{p.name} - Microsoft Edge"; edge_title2 = f"{p.stem} - Microsoft Edge"
                    app_tokens_arr = ["sarvajña", "sarvajna", "sarvajnagpt", "sarvajna gpt"]
                    browser_tokens_arr = [edge_title1, edge_title2, p.name, p.stem, stem_space_arr, "microsoft edge", "edge", "google chrome", "chrome"]
                    if page_title:
                        browser_tokens_arr = [f"{page_title} - Microsoft Edge", page_title] + browser_tokens_arr
                    code_tokens_arr = [code_title1, code_title2, "visual studio code", "vs code", "vscode", "code"]
                    tri_snap = {"attempted": True}
                    if callable(layout_any_right_stack):  # type: ignore[truthy-bool]
                        laidg = layout_any_right_stack(app_tokens_arr, browser_tokens_arr, code_tokens_arr)  # type: ignore[misc]
                        tri_snap["generic_stack"] = laidg
                        snap = {"attempted": True, "selected": bool(isinstance(laidg, dict) and laidg.get("ok")), "tri": True, "method": "generic_right_stack"}
                    else:
                        snap = {"attempted": True, "selected": False, "tri": True, "error": "generic_stack_unavailable"}
                except Exception:
                    snap = {"attempted": True, "selected": False, "tri": True, "error": "generic_stack_failed"}
            else:
                side = (req.side or 'right')
                snap = snap_current_and_select(tokens, snap_side=side)  # type: ignore[misc]
            try:
                if not snap.get("selected") and not is_code_like and is_word_like:
                    import time as _t
                    _t.sleep(0.12)
                # For non-Word flows, keep a minimal browser fallback
                if ((is_code_like or not is_word_like) and (not snap.get("selected")) and isinstance(snap.get("diagnostics"), dict)):
                    diag = snap.get("diagnostics") or {}
                    names = [str(n).lower() for n in (diag.get("unique_names") or [])]
                    def _has_sarvaj_browser(n: str) -> bool:
                        return (
                            ("sarvajna" in n or "sarvajña" in n)
                            and ("chrome" in n or "edge" in n or "firefox" in n or "brave" in n)
                        )
                    has_sarvaj_browser = any(_has_sarvaj_browser(n) for n in names)
                    if has_sarvaj_browser and callable(select_snap_assist_tile):
                        ok_br = bool(select_snap_assist_tile(["sarvajña", "sarvajna", "microsoft edge", "google chrome", "mozilla firefox", "brave"]))  # type: ignore[misc]
                        if ok_br:
                            snap["selected"] = True
                            snap["fallback"] = "browser_sarvaj_selected"
            except Exception:
                pass
        elif callable(select_snap_assist_tile):  # type: ignore[truthy-bool]
            sel = bool(select_snap_assist_tile(tokens))  # type: ignore[misc]
            snap = {"attempted": True, "selected": sel, "tokens": tokens}
        else:
            snap = {"attempted": False, "selected": False, "reason": "no_selector"}
    result = {
        "opened": launched,
        "path": str(p),
        # Consider split_screen true only if selection happened and verification passed
        "split_screen": bool(snap and (snap.get("selected") and (snap.get("verified") or snap.get("snap_sent")))),
        "cua_attempted": bool(snap and (snap.get("attempted") or snap.get("snap_sent"))),
        "cua_selected": bool(snap and (snap.get("selected") is True)),
        "snap": snap or {"attempted": False, "selected": False},
        "flow": "cua_only",
        "side": (req.side or 'right'),
    }
    if browser_info is not None:
        result["browser"] = browser_info
    if tri_snap is not None:
        result["tri_snap"] = tri_snap
    return result


def _cua_status() -> Dict[str, Any]:
    if callable(cua_runtime_status):  # type: ignore[truthy-bool]
        try:
            st = cua_runtime_status()  # type: ignore[misc]
            if isinstance(st, dict):
                return st
        except Exception:
            pass
    return {"available": False, "detail": "cua_runtime_status unavailable"}


def _select_tokens(tokens: list[str]) -> Dict[str, Any]:
    if not callable(select_snap_assist_tile):  # type: ignore[truthy-bool]
        return {"attempted": False, "selected": False, "reason": "select_snap_assist_tile_unavailable"}
    try:
        ok = bool(select_snap_assist_tile(tokens))  # type: ignore[misc]
        return {"attempted": True, "selected": ok, "tokens": tokens}
    except Exception as e:  # pragma: no cover
        return {"attempted": True, "selected": False, "tokens": tokens, "error": str(e)}

@router.get("/cua/status")
def power_cua_status() -> Dict[str, Any]:
    return {"ok": True, "status": _cua_status()}

@router.post("/open_doc_cua_only")
def open_doc_cua_only(req: OpenDocCUARequest) -> Dict[str, Any]:
    p = Path(req.path).expanduser()
    exists = p.is_file()
    launched = False
    if not exists:
        raise HTTPException(status_code=400, detail="path not found")
    # Use adapter helper for consistent behavior
    try:
        # For Word, prefer background launch so our app stays active
        suffix = p.suffix.lower()
        is_word_like = suffix in {".doc", ".docx", ".rtf"}
        if is_word_like and callable(cua_open_path_background):
            r = cua_open_path_background(str(p))  # type: ignore[misc]
            if not r.get("ok") and callable(cua_open_path):
                r = cua_open_path(str(p))
        else:
            r = cua_open_path(str(p)) if callable(cua_open_path) else {"ok": False, "error": "helper_missing"}
        launched = bool(r.get("ok"))
        if not launched and r.get("error"):
            raise HTTPException(status_code=500, detail=str(r.get("error")))
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"launch_failed: {e}")
    snap = None
    if bool(req.snap):
        stem = p.stem
        # Wait for Word window to appear and focus it
        try:
            if callable(wait_for_window_appearance):
                _ = wait_for_window_appearance([stem, "microsoft word", "word"], timeout_ms=2000)
            else:
                _ = None
        except Exception:
            _ = None
        try:
            if callable(focus_window_by_tokens):
                focus_window_by_tokens([stem, "microsoft word", "word"])  # best-effort focus on Word
            elif callable(wait_for_focus):
                wait_for_focus([stem, "microsoft word", "word"], timeout_ms=800)
        except Exception:
            pass
        try:
            import time as _t
            _t.sleep(0.08)
        except Exception:
            pass
        # Snap Word LEFT and select SarvajñaGPT tab on RIGHT; fallback to generic browser-only tokens if needed
        attempt_sets: list[list[str]] = [
            ["sarvajnagpt", "microsoft edge"], ["sarvajna gpt", "microsoft edge"],
            ["sarvajña", "microsoft edge"], ["sarvajna", "microsoft edge"],
            ["sarvajnagpt", "google chrome"], ["sarvajna gpt", "google chrome"],
            ["sarvajña", "google chrome"], ["sarvajna", "google chrome"],
            ["sarvajnagpt", "mozilla firefox"], ["sarvajna gpt", "mozilla firefox"],
            ["sarvajña", "mozilla firefox"], ["sarvajna", "mozilla firefox"],
            ["sarvajnagpt", "brave"], ["sarvajna gpt", "brave"],
            ["sarvajña", "brave"], ["sarvajna", "brave"],
        ]
        attempt_sets_generic: list[list[str]] = [["microsoft edge"], ["google chrome"], ["mozilla firefox"], ["brave"]]
        if callable(snap_current_and_select):  # type: ignore[truthy-bool]
            import time as _tlim
            snap = {"attempted": True, "selected": False}
            start = _tlim.time()
            # Trigger snap once using the first tokens
            snap = snap_current_and_select(attempt_sets[0], snap_side='left')  # type: ignore[misc]
            if not snap.get("selected"):
                # Within the remaining 2s window, try alternative tokens WITHOUT re-triggering snap
                for toks in attempt_sets[1:]:
                    if (_tlim.time() - start) > 2.0:
                        break
                    try:
                        ok = bool(select_snap_assist_tile(toks)) if callable(select_snap_assist_tile) else False
                        if ok:
                            snap["selected"] = True
                            snap["tokens"] = toks
                            break
                    except Exception:
                        break
            # Diagnostics-guided retry: if a tile containing SarvajñaGPT was seen, target it explicitly
            if not snap.get("selected") and isinstance(snap.get("diagnostics"), dict):
                try:
                    diag = snap.get("diagnostics") or {}
                    names = [str(n) for n in (diag.get("unique_names") or [])]
                    sar_cands = [n for n in names if ("sarvajna" in n.lower()) or ("sarvajña" in n.lower()) or ("sarvajnagpt" in n.lower())]
                    if sar_cands and callable(select_snap_assist_tile):
                        for cand in sar_cands[:4]:
                            try:
                                if bool(select_snap_assist_tile([cand])):
                                    snap["selected"] = True; snap["tokens"] = [cand]
                                    break
                                for brand in ("microsoft edge", "google chrome", "mozilla firefox", "brave"):
                                    if bool(select_snap_assist_tile([cand, brand])):
                                        snap["selected"] = True; snap["tokens"] = [cand, brand]
                                        break
                                if snap.get("selected"):
                                    break
                            except Exception:
                                continue
                except Exception:
                    pass
            # If still not selected, re-trigger once with generic browser tokens
            if not snap.get("selected"):
                try:
                    snap = snap_current_and_select(attempt_sets_generic[0], snap_side='left')  # type: ignore[misc]
                except Exception:
                    pass
                if not snap.get("selected"):
                    g_start = _tlim.time()
                    for toks in attempt_sets_generic[1:]:
                        if (_tlim.time() - g_start) > 1.8:
                            break
                        try:
                            ok = bool(select_snap_assist_tile(toks)) if callable(select_snap_assist_tile) else False
                            if ok:
                                snap["selected"] = True
                                snap["tokens"] = toks
                                break
                        except Exception:
                            break
        else:
            # Fallback single-shot using first attempt tokens
            snap = _select_tokens(attempt_sets[0])
    return {
        "ok": True,
        "launched": launched,
        "path": str(p),
        "snap": snap or {"attempted": False, "selected": False},
        "note": "CUA-only open; no Win32/COM/UIAutomation involved.",
    }

__all__ = [
    "router",
    "OpenDocCUARequest",
    "open_doc_cua_only",
    "power_cua_status",
    # Compatibility exports expected by main.py
    "OpenDocIntelligentlyRequest",
    "open_doc_intelligently",
    "compute_open_doc_signature",
]

# Optional: VS Code opener (pure CUA helper)
class OpenCodeCUARequest(BaseModel):
    path: Optional[str] = Field(None, description="Path to open in Code (optional)")
    new_window: bool | None = Field(True, description="Open in a new window (forced true by default)")

@router.post("/open_code_cua_only")
def open_code_cua_only(req: OpenCodeCUARequest) -> Dict[str, Any]:
    try:
        # Force new window even if an existing Code window is present
        result = cua_open_vscode(req.path, True) if callable(cua_open_vscode) else {"ok": False, "error": "helper_missing"}
        if not result.get("ok"):
            raise HTTPException(status_code=500, detail=str(result.get("error")))
        return {"ok": True, **result}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"open_vscode_failed: {e}")


# ---------------- Reopen tri-split for existing HTML artifacts ----------------
class ReopenTriSplitRequest(BaseModel):
    chat_id: str = Field(..., description="Chat ID to reopen tri-split for")
    service: Optional[str] = Field(None, description="Service name; defaults to 'power_mode'")


def _find_html_artifact_for_chat(chat_id: str, service: Optional[str]) -> Optional[str]:
    """Return best HTML path for a chat via chat_state.doc_path or latest chat_artifact .html."""
    DBP = os.path.join(os.path.dirname(__file__), 'chat_embeddings.db')
    try:
        conn = _sqlite3.connect(DBP)
        try:
            c = conn.cursor()
            svc = service or 'power_mode'
            # 1) chat_state.doc_path
            try:
                c.execute('''CREATE TABLE IF NOT EXISTS chat_state (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    chat_id TEXT,
                    service TEXT,
                    persistent_tags TEXT,
                    doc_path TEXT,
                    created_at INTEGER,
                    updated_at INTEGER
                )''')
            except Exception:
                pass
            c.execute('SELECT doc_path FROM chat_state WHERE chat_id=? AND service=? ORDER BY id DESC LIMIT 1', (chat_id, svc))
            row = c.fetchone()
            if row and row[0]:
                p = str(row[0])
                if p.lower().endswith('.html') and os.path.isfile(p):
                    return p
            # 2) Latest chat_artifact .html
            try:
                c.execute('''CREATE TABLE IF NOT EXISTS chat_artifact (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    chat_id TEXT,
                    service TEXT,
                    path TEXT,
                    created_at INTEGER
                )''')
            except Exception:
                pass
            c.execute('SELECT path FROM chat_artifact WHERE chat_id=? AND service=? ORDER BY id DESC', (chat_id, svc))
            rows = c.fetchall() or []
            for r in rows:
                try:
                    p = str(r[0] or '')
                    if p.lower().endswith('.html') and os.path.isfile(p):
                        return p
                except Exception:
                    continue
        finally:
            conn.close()
    except Exception:
        return None
    return None


@router.post('/reopen_trisplit_for_chat')
def reopen_trisplit_for_chat(req: ReopenTriSplitRequest) -> Dict[str, Any]:
    """Reopen the tri-split layout for a previous chat that produced an HTML artifact.

    Behavior:
    - Look up HTML path via chat_state.doc_path, fallback to chat_artifact .html
    - Open in VS Code (new window)
    - Open Edge/Browser preview (new window)
    - Attempt tri-split: App | Browser (preview) | Code
    """
    html_path = _find_html_artifact_for_chat(req.chat_id, req.service)
    if not html_path:
        raise HTTPException(status_code=404, detail='no_html_artifact_found')
    target = Path(html_path)
    if not target.is_file():
        raise HTTPException(status_code=404, detail='html_path_missing')
    opened = None
    try:
        opened = cua_open_vscode(str(target), True) if callable(cua_open_vscode) else {"ok": False}
    except Exception:
        opened = {"ok": False, "error": "vscode_launch_failed"}
    browser = None
    tri_snap = None
    try:
        if callable(cua_open_browser_to_path):
            browser = cua_open_browser_to_path(str(target), new_window=True)
            try:
                if callable(wait_for_window_appearance):
                    stem = target.stem
                    stem_space = stem.replace("-", " ")
                    _ = wait_for_window_appearance([target.name, stem, stem_space, "microsoft edge", "edge", "google chrome", "chrome", "mozilla firefox", "firefox", "brave"], timeout_ms=7000)
            except Exception:
                pass
        # Generic, order-agnostic right stack per request
        try:
            from .cua_adapter import layout_any_right_stack  # type: ignore
        except Exception:
            try:
                from cua_adapter import layout_any_right_stack  # type: ignore
            except Exception:
                layout_any_right_stack = None  # type: ignore
        try:
            stem = target.stem
            stem_space = stem.replace("-", " ")
            page_title = _html_title_from_file(target)
            code_title1 = f"{target.name} - Visual Studio Code"; code_title2 = f"{stem} - Visual Studio Code"
            edge_title1 = f"{target.name} - Microsoft Edge"; edge_title2 = f"{stem} - Microsoft Edge"
            app_tokens_arr = ["sarvajña", "sarvajna", "sarvajnagpt", "sarvajna gpt"]
            browser_tokens_arr = [edge_title1, edge_title2, target.name, stem, stem_space, "microsoft edge", "edge", "google chrome", "chrome"]
            if page_title:
                browser_tokens_arr = [f"{page_title} - Microsoft Edge", page_title] + browser_tokens_arr
            code_tokens_arr = [code_title1, code_title2, "visual studio code", "vs code", "vscode", "code"]
            if callable(layout_any_right_stack):  # type: ignore[truthy-bool]
                tri_snap = {"attempted": True, "generic_stack": layout_any_right_stack(app_tokens_arr, browser_tokens_arr, code_tokens_arr)}  # type: ignore[misc]
            else:
                tri_snap = {"attempted": True, "generic_stack": {"ok": False, "reason": "generic_stack_unavailable"}}
        except Exception:
            tri_snap = {"attempted": True, "generic_stack": {"ok": False, "reason": "generic_stack_failed"}}
    except Exception:
        browser = {"ok": False, "error": "browser_launch_failed"}
    return {
        "ok": True,
        "path": str(target),
        "opened": opened,
        **({"browser": browser} if browser else {}),
        **({"tri_snap": tri_snap} if tri_snap else {}),
    }


# ------------------------ Minimal power chat endpoint ------------------------
class PowerChatMessage(BaseModel):
    role: str = Field(..., description="user|assistant|system")
    content: str = Field(..., description="Message content")


class PowerChatRequest(BaseModel):
    # Flexible schema: accept either chat-style messages[] OR direct text fields from UI
    request_id: Optional[str] = Field(None, description="Optional request id")
    messages: list[PowerChatMessage] = Field(default_factory=list)
    mode: Optional[str] = Field(None, description="Optional mode, e.g., power_mode")
    # Alternate shape used by UI
    chat_id: Optional[str] = None
    text: Optional[str] = None
    doc_info: Optional[str] = None
    service: Optional[str] = None
    mem_tags: Optional[list[str]] = None
    selected_tags: Optional[list[str]] = None
    mem_context: Optional[str] = None
    auto_execute: Optional[bool] = Field(False, description="If true, server may plan/execute actions and include results")


# -------------------- Word inline enhance/replace endpoint --------------------
class WordEnhanceSelectionRequest(BaseModel):
    prompt: str = Field(..., description="Instruction for the LLM, e.g., 'explain in more detail'.")
    selection_text: str = Field(..., description="The selected text to enhance (from frontend clipboard)")
    chat_id: str | None = Field(None, description="Chat ID to lookup associated document path from chat_state")
    # Optional tuning
    max_full_context_chars: int | None = Field(60000, description="Cap full-document text for LLM context")
    rich: bool | None = Field(True, description="Paste with basic rich text formatting from Markdown output")


@router.post("/word/enhance_selection")
def word_enhance_selection(req: WordEnhanceSelectionRequest) -> Dict[str, Any]:
    """Enhance Word selection using provided text, send to LLM with full-document context, and replace selection.

    Flow:
    - Use selection text provided by frontend (already captured from clipboard)
    - If chat_id provided, look up Word document path from chat_state database
    - If doc_path found, read full document content directly from file
    - Otherwise fall back to Ctrl+A/Ctrl+C method to capture from open Word window
    - Call LLM with (prompt + selected + full context)
    - Paste replacement back into Word, replacing the selection
    Returns diagnostics and paste status.
    """
    import time as _t
    
    print("\n" + "="*80)
    print("WORD ENHANCEMENT DEBUG - START")
    print("="*80)
    
    # 1) Use the selection text provided by frontend
    selected_preview = (req.selection_text or '').strip()
    print(f"\n1. SELECTION TEXT RECEIVED:")
    print(f"   Length: {len(selected_preview)}")
    print(f"   Preview: {selected_preview[:200]}")
    print(f"   Chat ID: {req.chat_id}")
    print(f"   Prompt: {req.prompt}")
    
    if not selected_preview:
        raise HTTPException(status_code=400, detail="no_selection_provided - selection_text is empty")
    
    # 2) Try to get full document from database if chat_id provided
    full_text = ''
    doc_source = 'none'
    
    print(f"\n2. DOCUMENT LOOKUP:")
    if req.chat_id:
        print(f"   Chat ID provided: {req.chat_id}")
        # Look up document path from chat_state
        DBP = os.path.join(os.path.dirname(__file__), 'chat_embeddings.db')
        print(f"   Database path: {DBP}")
        try:
            conn = _sqlite3.connect(DBP)
            try:
                c = conn.cursor()
                c.execute('''CREATE TABLE IF NOT EXISTS chat_state (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    chat_id TEXT,
                    service TEXT,
                    persistent_tags TEXT,
                    doc_path TEXT,
                    created_at INTEGER,
                    updated_at INTEGER
                )''')
                # Look for doc_path in chat_state (service defaults to 'power_mode')
                c.execute('SELECT doc_path FROM chat_state WHERE chat_id=? ORDER BY id DESC LIMIT 1', (req.chat_id,))
                row = c.fetchone()
                print(f"   Database query result: {row}")
                if row and row[0]:
                    doc_path = str(row[0])
                    print(f"   Found doc_path: {doc_path}")
                    print(f"   File exists: {os.path.isfile(doc_path)}")
                    print(f"   Is Word file: {doc_path.lower().endswith(('.doc', '.docx'))}")
                    # Check if it's a Word document
                    if doc_path.lower().endswith(('.doc', '.docx')) and os.path.isfile(doc_path):
                        # Read the Word document
                        try:
                            from docx import Document  # type: ignore
                            doc = Document(doc_path)
                            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                            full_text = '\n'.join(paragraphs)
                            doc_source = f'database:{os.path.basename(doc_path)}'
                            print(f"   Successfully read from file!")
                            print(f"   Paragraphs found: {len(paragraphs)}")
                            print(f"   Total chars: {len(full_text)}")
                        except ImportError as e:
                            print(f"   ERROR: python-docx not available: {e}")
                            # python-docx not available, fall back to clipboard method
                            pass
                        except Exception as e:
                            print(f"   ERROR reading file: {e}")
                            # Error reading file, fall back to clipboard method
                            pass
                else:
                    print(f"   No doc_path found in database")
            finally:
                conn.close()
        except Exception as e:
            print(f"   Database error: {e}")
            pass
    else:
        print(f"   No chat_id provided, skipping database lookup")
    
    # 3) If no full text from database, use clipboard method (focus Word and capture)
    restored = False
    print(f"\n3. FULL DOCUMENT CAPTURE:")
    print(f"   Doc source so far: {doc_source}")
    print(f"   Full text length: {len(full_text)}")
    
    if not full_text:
        print(f"   No text from database, using clipboard method...")
        # Ensure Word is focused
        focused = False
        try:
            if callable(focus_window_by_tokens):
                focused = focus_window_by_tokens(["microsoft word", "word"])  # type: ignore[misc]
                print(f"   focus_window_by_tokens result: {focused}")
            if not focused and callable(ensure_focus):
                focused = ensure_focus(["microsoft word", "word"])  # type: ignore[misc]
                print(f"   ensure_focus result: {focused}")
            _t.sleep(0.15)  # Let focus settle
        except Exception as e:
            print(f"   Focus error: {e}")
            pass  # Continue anyway
        
        # Capture full-document text and restore selection
        try:
            if callable(capture_full_document_text_and_restore_selection):
                print(f"   Calling capture_full_document_text_and_restore_selection...")
                caps = capture_full_document_text_and_restore_selection(selected_preview, max_chars=(req.max_full_context_chars or 60000))  # type: ignore[misc]
                print(f"   Capture result: {caps}")
                if isinstance(caps, dict):
                    full_text = str(caps.get('text') or '')
                    restored = bool(caps.get('restored'))
                    doc_source = 'clipboard'
                    print(f"   Clipboard capture successful!")
                    print(f"   Full text length: {len(full_text)}")
                    print(f"   Selection restored: {restored}")
        except Exception as e:
            print(f"   Clipboard capture error: {e}")
            full_text = ''
            restored = False
    else:
        print(f"   Using text from database, skipping clipboard capture")
    
    # 4) Build LLM prompt with clear instructions
    user_prompt = (req.prompt or '').strip()
    if not user_prompt:
        user_prompt = 'Rewrite in more detail while preserving meaning and style.'
    
    # Keep context manageable
    max_ctx = max(10000, int(req.max_full_context_chars or 60000))
    ctx = (full_text or '')
    if len(ctx) > max_ctx:
        ctx = ctx[:max_ctx]
    
    print(f"\n4. LLM PROMPT CONSTRUCTION:")
    print(f"   User prompt: {user_prompt}")
    print(f"   Context length (before truncation): {len(full_text)}")
    print(f"   Context length (after truncation): {len(ctx)}")
    print(f"   Max context allowed: {max_ctx}")
    
    # Enhanced LLM prompt - Just enhance the selected text
    llm_input = f"""You are a text enhancement assistant. Your ONLY job is to improve the selected text based on the user's request.

CRITICAL RULES:
- Return ONLY the improved text, nothing else
- NO explanations, NO thinking process, NO commentary
- NO markers like "Here is..." or "ENHANCED TEXT:"
- Just the improved text that will replace the selection

USER REQUEST: {user_prompt}

SELECTED TEXT TO ENHANCE:
{selected_preview}

FULL DOCUMENT CONTEXT (for reference only):
{ctx}

Now return ONLY the enhanced version of the selected text:"""
    
    print(f"\n   FULL LLM INPUT:")
    print("   " + "-"*76)
    print("   " + llm_input.replace("\n", "\n   "))
    print("   " + "-"*76)
    
    # 5) Call LLM
    print(f"\n5. CALLING LLM...")
    try:
        out_text = _power_llm(llm_input) if callable(_power_llm) else ''  # type: ignore[misc]
        print(f"   LLM Response received!")
        print(f"   Response length: {len(out_text)}")
        print(f"   Response preview (first 500 chars):")
        print("   " + "-"*76)
        print("   " + out_text[:500].replace("\n", "\n   "))
        print("   " + "-"*76)
        
        # Clean up LLM response - remove <think> tags and extra commentary
        import re
        # Remove <think>...</think> blocks
        out_text = re.sub(r'<think>.*?</think>', '', out_text, flags=re.DOTALL | re.IGNORECASE)
        # Remove common prefixes
        out_text = re.sub(r'^(Here is|Here\'s|Enhanced text:|ENHANCED TEXT:)\s*', '', out_text, flags=re.IGNORECASE | re.MULTILINE)
        # Remove markdown code blocks if present
        out_text = re.sub(r'^```.*?\n', '', out_text, flags=re.MULTILINE)
        out_text = re.sub(r'\n```$', '', out_text)
        out_text = out_text.strip()
        
        print(f"   After cleanup:")
        print(f"   Cleaned length: {len(out_text)}")
        print(f"   Cleaned preview (first 500 chars):")
        print("   " + "-"*76)
        print("   " + out_text[:500].replace("\n", "\n   "))
        print("   " + "-"*76)
        
    except Exception as e:
        print(f"   LLM ERROR: {e}")
        raise HTTPException(status_code=500, detail=f"llm_failed:{e}")
    
    out_text = (out_text or '').strip()
    if not out_text:
        print(f"   ERROR: LLM returned empty output!")
        raise HTTPException(status_code=500, detail="llm_empty_output")
    
    print(f"   LLM output ready for pasting")
    
    # 6) Copy the enhanced text to clipboard for manual pasting
    print(f"\n6. COPYING TO CLIPBOARD:")
    try:
        import pyperclip  # type: ignore
        pyperclip.copy(out_text)
        print(f"   ✓ Enhanced text copied to clipboard ({len(out_text)} chars)")
        clipboard_copied = True
    except Exception as e:
        print(f"   ✗ Failed to copy to clipboard: {e}")
        clipboard_copied = False
    
    result = {
        'ok': True,
        'selection': {'length': len(selected_preview), 'preview': selected_preview[:240]},
        'context_used': len(ctx),
        'full_doc_length': len(full_text),
        'doc_source': doc_source,
        'restored_before_replace': bool(restored),
        'clipboard_copied': clipboard_copied,
        'llm_output_length': len(out_text),
        'enhanced_text_preview': out_text[:500],  # Preview for debugging
        'instructions': 'Select the text in Word → Press Ctrl+V to replace with enhanced version',
    }
    
    print(f"\n7. FINAL RESULT:")
    print(f"   {result}")
    print("\n" + "="*80)
    print("WORD ENHANCEMENT DEBUG - END")
    print("="*80 + "\n")
    
    return result


# Lightweight LLM + hashtag router for power chat
try:
    from .llm_inference import generate_response as _power_llm
except Exception:  # pragma: no cover
    try:
        from llm_inference import generate_response as _power_llm  # type: ignore
    except Exception:  # pragma: no cover
        def _power_llm(prompt: str) -> str:  # type: ignore
            return ""

try:
    # Import agent autoplan to support #plan hashtag
    from .agent_router import autoplan as _agent_autoplan, NaturalLanguagePlanRequest as _PlanReq
except Exception:  # pragma: no cover
    _agent_autoplan = None  # type: ignore
    _PlanReq = None  # type: ignore


@router.post("/power_chat")
def power_chat(req: PowerChatRequest) -> Dict[str, Any]:
    # Determine conversational content
    last = (req.text or '').strip()
    roles = []
    if req.messages:
        roles = [(m.role, m.content) for m in (req.messages or [])]
        if not last:
            last = (req.messages[-1].content or '').strip()
    else:
        if last:
            roles = [("user", last)]
    import re as _re
    tags_inline = _re.findall(r"#([A-Za-z0-9_\-]+)", last)
    # Merge inline, mem_tags, and selected_tags; normalize + dedupe preserving order
    merge_src = []
    merge_src.extend(tags_inline)
    if req.mem_tags:
        merge_src.extend([str(t) for t in req.mem_tags if t])
    if req.selected_tags:
        merge_src.extend([str(t) for t in req.selected_tags if t])
    seen = set()
    normalized_tags: list[str] = []
    for t in merge_src:
        k = str(t).strip().lstrip('#').lower()
        if k and k not in seen:
            seen.add(k)
            normalized_tags.append(k)

    # Helper: resolve memory context and quick item previews for given tag list
    def _mem_from_tags(tag_list: list[str]) -> tuple[Optional[str], list[Dict[str, Any]]]:
        if not tag_list:
            return None, []
        import sqlite3 as _sqlite3
        mem_ctx_parts: list[str] = []
        items: list[Dict[str, Any]] = []
        try:
            conn = _sqlite3.connect('chat_embeddings.db')
            c = conn.cursor()
            for t in tag_list:
                if not t:
                    continue
                like = f"%{t}%"
                try:
                    c.execute(
                        'SELECT id, substr(text_content,1,500) as preview, filename FROM mem_item '
                        'WHERE tags LIKE ? OR filename LIKE ? ORDER BY id DESC LIMIT 5',
                        (like, like)
                    )
                    rows = c.fetchall() or []
                except Exception:
                    rows = []
                if rows:
                    previews = [(r[1] or '') for r in rows]
                    mem_ctx_parts.append(f"Tag: #{t}\n" + '\n---\n'.join(previews))
                    for r in rows:
                        items.append({"id": r[0], "filename": r[2], "preview": r[1], "tag": t})
            try:
                conn.close()
            except Exception:
                pass
        except Exception:
            pass
        mem_ctx = '\n\n'.join(mem_ctx_parts) if mem_ctx_parts else None
        # Trim to a reasonable size to avoid over-long prompts
        if mem_ctx and len(mem_ctx) > 4000:
            mem_ctx = mem_ctx[:4000] + "\n..."
        return mem_ctx, items

    # If client already provided mem_context, honor it; else resolve from tags
    provided_mem_ctx = (req.mem_context or '').strip() if isinstance(getattr(req, 'mem_context', None), str) else None

    # Route: #plan triggers agent autoplan (no execute)
    if "plan" in normalized_tags and _agent_autoplan and _PlanReq:
        try:
            # Remove hashtags from the prompt sent to the planner
            plain = _re.sub(r"#[A-Za-z0-9_\-]+", "", last).strip()
            # If other memory tags are present, resolve and prepend as context
            mem_tags_only = [t for t in normalized_tags if t != 'plan']
            mem_ctx = provided_mem_ctx
            mem_items: list[Dict[str, Any]] = []
            if not mem_ctx and mem_tags_only:
                mem_ctx, mem_items = _mem_from_tags(mem_tags_only)
            if mem_ctx:
                plain = ("Memory context from #tags:\n" + mem_ctx + "\n\nTask: " + (plain or last or "")).strip()
            plan_req = _PlanReq(prompt=plain or last or "Generate a simple plan", execute=False, base_folder="project")  # type: ignore
            plan_resp = _agent_autoplan(plan_req)  # returns pydantic model
            return {
                "ok": True,
                "request_id": req.request_id,
                "mode": req.mode or "power_mode",
                "tags": normalized_tags,
                "plan": {
                    "raw": plan_resp.raw_plan,
                    "actions": [a.dict() for a in plan_resp.actions],  # type: ignore
                    "executed": plan_resp.executed,
                },
                **({"memory": {"used_tags": mem_tags_only, "items": mem_items}} if mem_ctx else {}),
                # Frontend often prefers assistant_text; provide both keys
                "assistant_text": "Created a plan based on your request" if plan_resp.actions else "No actions planned.",
            }
        except Exception as e:
            # Fall through to LLM text if planning fails
            pass

    # Default: simple LLM echo of the conversation
    try:
        # Build a lightweight chat-style prompt
        history = []
        for r, c in roles[-6:]:
            history.append(f"{r}: {c}")
        # Resolve memory context (if any tags were provided or passed in)
        mem_items: list[Dict[str, Any]] = []
        mem_ctx = provided_mem_ctx
        if not mem_ctx and normalized_tags:
            mem_ctx, mem_items = _mem_from_tags(normalized_tags)
        base = "You are Sarvajña Power Mode. Be concise and helpful.\n"
        if mem_ctx:
            base += "Memory context from #tags:\n" + mem_ctx + "\n\n"
        # If the intent is code, enforce code-only output with a single fenced block
        low_text = (last or '').lower()
        intent_code_prompt = ("code" in low_text or "vscode" in low_text or "visual studio code" in low_text or "html" in low_text or "python" in low_text or "javascript" in low_text or "typescript" in low_text)
        if intent_code_prompt:
            base += (
                "When the user asks for code, output ONLY a single fenced code block for the requested language.\n"
                "Do not include explanations, headings, notes, or <think> sections.\n"
                "Improve the code by self-review internally, but only return the final code in one fenced block.\n"
                "Prefer complete, production-ready structure (e.g., full HTML skeleton for HTML).\n"
            )
            # If HTML is involved, raise the bar: responsiveness, animation, and design quality
            if 'html' in (last or '').lower():
                base += (
                    "For HTML/CSS tasks:\n"
                    "- Include <!doctype html>, <meta viewport>, and semantic structure.\n"
                    "- Use modern responsive layout (CSS Grid/Flexbox), fluid typography, and CSS variables for a cohesive color system.\n"
                    "- Apply tasteful, accessible color palettes (adequate contrast), and system fonts or Google Fonts when appropriate.\n"
                    "- Add subtle animations/transitions (e.g., fade/slide, hover effects) that feel smooth and performant (prefers-reduced-motion respected).\n"
                    "- Ensure mobile-first responsiveness; verify sections stack gracefully.\n"
                    "- Keep JS minimal and unobtrusive (only if needed).\n"
                    "Internally iterate 2x to refine layout, hierarchy, and polish; output only the final single fenced HTML code block.\n"
                )
        # Include current user instruction explicitly if messages[] were not provided
        user_line = f"user: {last}" if (last and not req.messages) else ""
        prompt = base + "\n".join(history + ([user_line] if user_line else [])) + "\nassistant:"
        text = _power_llm(prompt)
    except Exception:
        text = ""
    # Helper: default save directory under Documents/Sarvjan
    def _default_save_dir() -> Path:
        try:
            home = Path.home()
            docs = home / 'Documents'
            target = docs / 'SarvajnaGPT'
            target.mkdir(parents=True, exist_ok=True)
            return target
        except Exception:
            # Fallback to repo-root folder if Documents isn't available
            try:
                repo_root = Path(__file__).resolve().parent.parent
                target = repo_root / 'SarvajnaGPT'
                target.mkdir(parents=True, exist_ok=True)
                return target
            except Exception:
                return Path.cwd()

    # Optional auto-execute: if the user asked to open Word and write, or open VS Code and write code, act accordingly
    execute_block: Dict[str, Any] | None = None
    try:
        want_exec = bool(getattr(req, 'auto_execute', False))
        low = (last or '').lower()
        intent_word_write = ("word" in low or "document" in low) and ("write" in low or "type" in low or "insert" in low or "fill" in low or "paste" in low or "create" in low or "save" in low or "summary" in low)
        intent_code = ("code" in low or "vscode" in low or "visual studio code" in low) and ("write" in low or "create" in low or "make" in low or "generate" in low)

        if want_exec and intent_code and text:
            # Create a new file under Documents/Sarvjan and open in VS Code (new window)
            save_dir = _default_save_dir()
            # Basic extension heuristics
            ext = '.txt'
            lang_hint = None
            if any(k in low for k in ['python', 'py code']):
                ext = '.py'
                lang_hint = 'python'
            elif any(k in low for k in ['javascript', 'js code', 'node']):
                ext = '.js'
                lang_hint = 'javascript'
            elif 'typescript' in low or 'ts code' in low:
                ext = '.ts'
                lang_hint = 'typescript'
            elif 'html' in low:
                ext = '.html'
                lang_hint = 'html'
            elif 'css' in low:
                ext = '.css'
                lang_hint = 'css'
            if not lang_hint:
                lang_hint = _guess_code_language(text)
            ts = _dt.datetime.now().strftime('%Y%m%d-%H%M%S')
            fname = f"code-{ts}{ext}"
            target = save_dir / fname
            # Write assistant text to file
            try:
                code_only = _extract_code_from_llm(text or '', lang_hint)
                target.write_text(code_only or '', encoding='utf-8')
            except Exception as e:
                # Ensure directory exists, then retry once
                try:
                    target.parent.mkdir(parents=True, exist_ok=True)
                    code_only = _extract_code_from_llm(text or '', lang_hint)
                    target.write_text(code_only or '', encoding='utf-8')
                except Exception:
                    pass
            opened = None
            try:
                # Always open in a NEW VS Code window for deterministic snapping
                opened = cua_open_vscode(str(target), True) if callable(cua_open_vscode) else {"ok": False}
            except Exception:
                opened = {"ok": False, "error": "vscode_launch_failed"}
            # If HTML, open a browser preview of the file (prefer Chrome) in its own window (avoid duplicate tabs)
            browser = None
            tri_snap = None
            if ext == '.html' and callable(cua_open_browser_to_path):
                try:
                    browser = cua_open_browser_to_path(str(target), new_window=True)
                    # Order-agnostic right-stack flow per user request
                    try:
                        from .cua_adapter import layout_any_right_stack  # type: ignore
                    except Exception:
                        try:
                            from cua_adapter import layout_any_right_stack  # type: ignore
                        except Exception:
                            layout_any_right_stack = None  # type: ignore
                    try:
                        stem = target.stem
                        stem_space = stem.replace("-", " ")
                        page_title = _html_title_from_file(target)
                        if callable(wait_for_window_appearance):
                            wait_tokens = [target.name, stem, stem_space, "microsoft edge", "edge", "google chrome", "chrome", "mozilla firefox", "firefox", "brave"]
                            if page_title:
                                wait_tokens = [page_title] + wait_tokens
                            _ = wait_for_window_appearance(wait_tokens, timeout_ms=7000)
                        code_title1 = f"{target.name} - Visual Studio Code"; code_title2 = f"{stem} - Visual Studio Code"
                        edge_title1 = f"{target.name} - Microsoft Edge"; edge_title2 = f"{stem} - Microsoft Edge"
                        app_tokens_arr = ["sarvajña", "sarvajna", "sarvajnagpt", "sarvajna gpt"]
                        browser_tokens_arr = [edge_title1, edge_title2, target.name, stem, stem_space, "microsoft edge", "edge", "google chrome", "chrome"]
                        if page_title:
                            browser_tokens_arr = [f"{page_title} - Microsoft Edge", page_title] + browser_tokens_arr
                        code_tokens_arr = [code_title1, code_title2, "visual studio code", "vs code", "vscode", "code"]
                        if callable(layout_any_right_stack):  # type: ignore[truthy-bool]
                            tri_snap = {"attempted": True, "generic_stack": layout_any_right_stack(app_tokens_arr, browser_tokens_arr, code_tokens_arr)}  # type: ignore[misc]
                        else:
                            tri_snap = {"attempted": True, "generic_stack": {"ok": False, "reason": "generic_stack_unavailable"}}
                    except Exception:
                        tri_snap = {"attempted": True, "generic_stack": {"ok": False, "reason": "generic_stack_failed"}}
                except Exception:
                    browser = {"ok": False, "error": "browser_launch_failed"}
            execute_block = {
                "intent": "code_write",
                "path": str(target),
                "opened": opened,
                **({"browser": browser} if browser else {}),
                **({"tri_snap": tri_snap} if tri_snap else {}),
            }

        if want_exec and intent_word_write and text and execute_block is None:
            # Find a target .doc or .docx file from memory items or doc_info
            from pathlib import Path as _P
            repo_root = _P(__file__).resolve().parent.parent
            search_dirs = [repo_root / 'uploads', repo_root / 'memory_uploads']

            def _find_file(fname: str) -> Optional[str]:
                if not fname:
                    return None
                for d in search_dirs:
                    try:
                        if d.is_dir():
                            cand = d / fname
                            if cand.is_file():
                                return str(cand)
                            # also search recursively as fallback (limit depth)
                            for p in d.rglob(fname):
                                try:
                                    if p.is_file() and p.name.lower() == fname.lower():
                                        return str(p)
                                except Exception:
                                    pass
                    except Exception:
                        continue
                return None

            def _candidate_from_mem() -> Optional[str]:
                try:
                    for it in (mem_items or []):
                        fn = str((it.get('filename') or '')).strip()
                        if fn and (fn.lower().endswith('.docx') or fn.lower().endswith('.doc')):
                            p = _find_file(fn)
                            if p:
                                return p
                except Exception:
                    pass
                return None

            target_path = None
            # 1) Prefer memory-listed filenames
            target_path = _candidate_from_mem()
            # 2) If doc_info mentions a filename, try that
            if not target_path and getattr(req, 'doc_info', None):
                try:
                    import re as _re2
                    m = _re2.search(r"([A-Za-z0-9_\-\s]+\.(?:docx|doc))", str(req.doc_info))
                    if m:
                        target_path = _find_file(m.group(1))
                except Exception:
                    pass
            # 3) If user text contains a filename
            if not target_path:
                try:
                    import re as _re3
                    m = _re3.search(r"([A-Za-z0-9_\-\s]+\.(?:docx|doc))", last or '')
                    if m:
                        target_path = _find_file(m.group(1))
                except Exception:
                    pass

            opened = None
            snap_info = None
            paste_info = None
            word_tokens: list[str] = ["microsoft word", "word"]
            # If we have a specific file, open and snap left; else best-effort focus existing Word and paste
            if target_path:
                try:
                    req_open = OpenDocIntelligentlyRequest(abs_path=target_path, split_screen=True, side='left', snap=True)
                    opened = open_doc_intelligently(req_open)
                    snap_info = opened.get('snap') if isinstance(opened, dict) else None
                except Exception:
                    opened = {"opened": False, "error": "open_failed"}
                # Focus Word window by doc name if possible, then paste
                try:
                    stem = _P(target_path).stem
                    toks = [stem] + word_tokens
                    cleaned = _plain_from_markdown(text)
                    rtf = _rtf_from_markdown(text)
                    if callable(paste_rich_text_to_foreground_app):
                        paste_info = paste_rich_text_to_foreground_app(cleaned, rtf, click_center=True, focus_tokens=toks)
                    else:
                        paste_info = paste_text_to_foreground_app(cleaned, click_center=True, focus_tokens=toks) if callable(paste_text_to_foreground_app) else None
                except Exception:
                    paste_info = {"ok": False, "error": "paste_failed"}
            else:
                # No path found: create a new .docx under Documents/Sarvjan, write text, then open it in Word
                try:
                    save_dir = _default_save_dir()
                    ts = _dt.datetime.now().strftime('%Y%m%d-%H%M%S')
                    out_path = save_dir / f"document-{ts}.docx"
                    ok_docx = _docx_from_markdown(text or '', str(out_path))
                    if not ok_docx:
                        # Fallback: write as .txt if conversion failed
                        out_path = save_dir / f"document-{ts}.txt"
                        out_path.write_text(_plain_from_markdown(text or ''), encoding='utf-8')
                    # Open created doc and snap
                    try:
                        req_open = OpenDocIntelligentlyRequest(abs_path=str(out_path), split_screen=True, side='left', snap=True)
                        opened = open_doc_intelligently(req_open)
                        snap_info = opened.get('snap') if isinstance(opened, dict) else None
                    except Exception:
                        opened = {"opened": False, "error": "open_failed"}
                    paste_info = None  # Not needed since content already saved in doc
                    execute_block = {
                        "intent": "word_write",
                        "path": str(out_path),
                        "opened": opened,
                        "snap": snap_info,
                        "paste": paste_info,
                    }
                except Exception:
                    # As a last resort, try to focus Word and paste into the active doc
                    try:
                        paste_info = paste_text_to_foreground_app(text, click_center=True, focus_tokens=word_tokens) if callable(paste_text_to_foreground_app) else None
                    except Exception:
                        paste_info = {"ok": False, "error": "paste_failed"}
                    execute_block = {
                        "intent": "word_write",
                        "opened": None,
                        "snap": None,
                        "paste": paste_info,
                    }
    except Exception:
        execute_block = None

    # Prepare assistant_text for chat: if we executed an action, summarize instead of dumping full document content
    chat_assistant_text = text
    try:
        if execute_block is not None and isinstance(execute_block, dict):
            intent = execute_block.get('intent')
            pth = execute_block.get('path') or (execute_block.get('opened') or {}).get('path') if isinstance(execute_block.get('opened'), dict) else None
            if intent == 'word_write':
                # Use cleaned text for count if available
                try:
                    words = len((_plain_from_markdown(text or '')).split())
                except Exception:
                    words = None
                if words:
                    chat_assistant_text = f"Inserted formatted content into Word ({words} words). Saved at: {pth or 'current document'}."
                else:
                    chat_assistant_text = f"Inserted formatted content into Word. Saved at: {pth or 'current document'}."
            elif intent == 'code_write':
                ext = ''
                try:
                    from pathlib import Path as _Pth
                    if pth:
                        ext = _Pth(pth).suffix
                except Exception:
                    ext = ''
                chat_assistant_text = f"Created a {ext or ''} file with your code at: {pth}. Opened in VS Code."
    except Exception:
        pass

    resp = {
        "ok": True,
        "request_id": req.request_id,
        "mode": req.mode or "power_mode",
        "tags": normalized_tags,
        "text": text,
        "assistant_text": chat_assistant_text,
        **({"memory": {"used_tags": normalized_tags, "items": mem_items}} if 'mem_items' in locals() and mem_items else {}),
    }
    if execute_block is not None:
        resp["execute"] = execute_block
        # Persist doc_path into chat_state if available
        try:
            doc_path = None
            if isinstance(execute_block, dict):
                # prefer explicit path if present
                doc_path = execute_block.get('path')
                # else see if opened result has path
                if not doc_path:
                    opened = execute_block.get('opened') or {}
                    if isinstance(opened, dict):
                        doc_path = opened.get('path')
            # Upsert chat_state with doc_path and tags
            if getattr(req, 'chat_id', None):
                DBP = os.path.join(os.path.dirname(__file__), 'chat_embeddings.db')
                conn_cs = _sqlite3.connect(DBP)
                try:
                    # Ensure table exists (best-effort)
                    ccs = conn_cs.cursor()
                    ccs.execute('''CREATE TABLE IF NOT EXISTS chat_state (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        chat_id TEXT,
                        service TEXT,
                        persistent_tags TEXT,
                        doc_path TEXT,
                        created_at INTEGER,
                        updated_at INTEGER
                    )''')
                    # Read existing
                    service = req.service or 'power_mode'
                    ccs.execute('SELECT id, persistent_tags, doc_path FROM chat_state WHERE chat_id=? AND service=? ORDER BY id DESC LIMIT 1', (req.chat_id, service))
                    row = ccs.fetchone()
                    # tags CSV from normalized_tags with leading '#'
                    tags_csv = None
                    try:
                        if normalized_tags:
                            tags_csv = ','.join([('#' + t) for t in normalized_tags if t])
                    except Exception:
                        tags_csv = None
                    nowi = int(_dt.datetime.now().timestamp())
                    if row:
                        # Merge: overwrite provided fields, keep missing
                        prev_tags = str(row[1] or '')
                        if tags_csv is None:
                            tags_csv = prev_tags or None
                        prev_doc = row[2]
                        if not doc_path:
                            doc_path = prev_doc
                        ccs.execute('UPDATE chat_state SET persistent_tags=?, doc_path=?, updated_at=? WHERE id=?', (tags_csv, doc_path, nowi, row[0]))
                    else:
                        ccs.execute('INSERT INTO chat_state (chat_id, service, persistent_tags, doc_path, created_at, updated_at) VALUES (?, ?, ?, ?, ?, ?)', (req.chat_id, service, tags_csv, doc_path, nowi, nowi))
                    conn_cs.commit()
                finally:
                    conn_cs.close()
        except Exception:
            pass

        # Record artifact for cleanup on chat deletion
        try:
            if getattr(req, 'chat_id', None):
                DBP = os.path.join(os.path.dirname(__file__), 'chat_embeddings.db')
                conn_af = _sqlite3.connect(DBP)
                try:
                    caf = conn_af.cursor()
                    caf.execute('''CREATE TABLE IF NOT EXISTS chat_artifact (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        chat_id TEXT,
                        service TEXT,
                        path TEXT,
                        created_at INTEGER
                    )''')
                    # pick best path candidate from execute
                    apath = None
                    if isinstance(execute_block, dict):
                        apath = execute_block.get('path')
                        if not apath:
                            opened = execute_block.get('opened') or {}
                            if isinstance(opened, dict):
                                apath = opened.get('path')
                    if apath:
                        nowi = int(_dt.datetime.now().timestamp())
                        caf.execute('INSERT INTO chat_artifact (chat_id, service, path, created_at) VALUES (?, ?, ?, ?)',
                                    (req.chat_id, (req.service or 'power_mode'), str(apath), nowi))
                        conn_af.commit()
                finally:
                    conn_af.close()
        except Exception:
            pass

    # Persist this turn into chat table (like /api/chat)
    try:
        if getattr(req, 'chat_id', None) and (last or text):
            DBP = os.path.join(os.path.dirname(__file__), 'chat_embeddings.db')
            conn_c = _sqlite3.connect(DBP)
            try:
                cc = conn_c.cursor()
                # Ensure chat table has expected columns
                cc.execute('''CREATE TABLE IF NOT EXISTS chat (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    chat_id TEXT,
                    user TEXT,
                    llm TEXT,
                    embedding BLOB,
                    timestamp INTEGER,
                    tags TEXT,
                    doc_info TEXT,
                    service TEXT,
                    chat_name TEXT
                )''')
                # Determine chat_name: reuse existing for chat_id else compute next 'New Chat N'
                cc.execute('SELECT DISTINCT chat_name FROM chat WHERE chat_id=? AND chat_name IS NOT NULL', (req.chat_id,))
                row = cc.fetchone()
                if row and row[0]:
                    chat_name = row[0]
                else:
                    cc.execute("SELECT DISTINCT chat_name FROM chat WHERE chat_name LIKE 'New Chat %'")
                    existing = [r[0] for r in cc.fetchall() if r and r[0]]
                    max_n = 0
                    for en in existing:
                        try:
                            n = int(str(en).replace('New Chat ', '').strip())
                            if n > max_n:
                                max_n = n
                        except Exception:
                            continue
                    chat_name = f'New Chat {max_n + 1}'
                # Compose tags CSV from normalized_tags
                tags_csv = None
                if normalized_tags:
                    try:
                        tags_csv = ','.join([('#' + t) for t in normalized_tags if t])
                    except Exception:
                        tags_csv = None
                # Insert row (store assistant_text shown to user, not the full document content)
                cc.execute('INSERT INTO chat (chat_id, user, llm, embedding, timestamp, doc_info, service, chat_name, tags) VALUES (?, ?, ?, ?, strftime("%s", "now"), ?, ?, ?, ?)',
                           (req.chat_id, last, chat_assistant_text, None, req.doc_info, (req.service or 'power_mode'), chat_name, tags_csv))
                conn_c.commit()
            finally:
                conn_c.close()
    except Exception:
        pass
    return resp
